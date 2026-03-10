from __future__ import annotations

from io import BytesIO
import json
import zipfile

from docx import Document

from kb_pipeline import KBDraft, TopicDocument, split_draft_into_topic_documents


def export_draft_to_docx_bytes(draft: KBDraft) -> bytes:
    document = Document()
    document.add_heading(draft.title, level=0)
    document.add_paragraph(draft.summary)

    if draft.warnings:
        document.add_heading("Warnings", level=1)
        for warning in draft.warnings:
            document.add_paragraph(warning, style="List Bullet")

    if draft.visual_notes:
        document.add_heading("Visual Notes", level=1)
        for note in draft.visual_notes:
            document.add_paragraph(note, style="List Bullet")

    if draft.table_notes:
        document.add_heading("Table Notes", level=1)
        for note in draft.table_notes:
            document.add_paragraph(note, style="List Bullet")

    for section in draft.sections:
        document.add_heading(section.heading, level=1)
        document.add_paragraph(section.content)
        if section.source_pages:
            document.add_paragraph(f"Source pages: {', '.join(map(str, section.source_pages))}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_topic_bundle_zip_bytes(draft: KBDraft) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("master-kb-draft.docx", export_draft_to_docx_bytes(draft))
        for topic in split_draft_into_topic_documents(draft):
            topic_doc = export_topic_document_bytes(topic)
            safe_name = topic.title.lower().replace(" ", "-")[:50] or topic.topic_id
            archive.writestr(f"{safe_name}.docx", topic_doc)
    return buffer.getvalue()


def export_topic_document_bytes(topic: TopicDocument) -> bytes:
    document = Document()
    document.add_heading(topic.title, level=0)
    document.add_paragraph(topic.summary)
    for section in topic.sections:
        document.add_heading(section.heading, level=1)
        document.add_paragraph(section.content)
    payload = BytesIO()
    document.save(payload)
    return payload.getvalue()


def export_share_package_bytes(draft: KBDraft, share_code: str, *, share_note: str = "", source_filename: str = "") -> bytes:
    buffer = BytesIO()
    topic_docs = split_draft_into_topic_documents(draft)
    manifest = {
        "share_code": share_code,
        "title": draft.title,
        "summary": draft.summary,
        "topic_count": len(topic_docs),
        "warnings": draft.warnings,
        "share_note": share_note,
        "source_filename": source_filename,
    }
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("share-manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
        archive.writestr("master-kb-draft.docx", export_draft_to_docx_bytes(draft))
        for topic in topic_docs:
            safe_name = topic.title.lower().replace(" ", "-")[:50] or topic.topic_id
            archive.writestr(f"topics/{safe_name}.docx", export_topic_document_bytes(topic))
    return buffer.getvalue()
